import os
import re
import cv2
import pyzipper
from typing import List
import numpy as np
import subprocess
import fitz
import datetime
import pandas as pd
import getpass
import docx
from pathlib import Path
import pdfplumber
from transformers import pipeline
import win32security
import win32api
from odf import text, teletype
from odf.opendocument import load
from natasha import (
    Segmenter,
    MorphVocab,
    NewsEmbedding,
    NewsNERTagger,
    Doc
)
from paddleocr import PaddleOCR
from PIL import Image, ImageDraw, ImageFilter
ocr = PaddleOCR(
    use_doc_orientation_classify=False,
    use_doc_unwarping=False,
    use_textline_orientation=False,
    lang="ru")

all_FIO = ["Фамилия:","Фамилия", "фамилия:","фамилия","Отчество:","Отчество","отчество:","отчество","Имя:","Имя","имя:","имя"]
segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
ner_tagger = NewsNERTagger(emb)


ner_pipeline = pipeline(
    "ner",
    model="rubert-pdn-final",
    tokenizer="rubert-pdn-final",
    aggregation_strategy="simple"
)

GENERAL_DATA_CLASSIFIER = {
    'EMAIL': 'Общие', 'Номер_Тел': 'Общие', 'Пасспорт': 'Общие', 'ИНН': 'Общие',
    'СНИЛС': 'Общие', 'Адрес': 'Общие', 'Полное_имя': 'Общие', 'Местоположение': 'Общие'
}
SPECIAL_DATA_MARKERS = {
    'Состояние здоровья': ['здоровье', 'диагноз', 'заболевание', 'медицинский', 'показания', 'противопоказания',
                           'анализы', 'госпитализация','тяжело болен','при смерти', 'рак'],
    'Политические взгляды': ['партия', 'политические', 'убеждения', 'состоит в', 'член партии', 'митинг', 'левые'],
    'Национальность': ['национальность', 'раса', 'этническое происхождение'],
    'Религиозные убеждения': ['религия', 'вероисповедание', 'конфессия', 'атеист', 'приход']
}


def init_destruction_log():
    if not os.path.exists(LOG_FILE):
        with open(LOG_FILE, "w", encoding="utf-8") as f:
            f.write(
                "Timestamp;User;FilePath;FileSize;Method;Passes;Result\n"
            )

def find_personal_data_in_text(text: str):
    all_findings = []
    patterns = {
        'EMAIL': r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        'Номер_Тел': r'(?:\+7|8)[\s\(-]*\d{3}[\s\)-]*\d{3}[\s-]*\d{2}[\s-]*\d{2}|(?:телефон[:\s]*)?(?:\+7|8)[\s\-()]*\d{3}[\s\-()]*\d{3}[\s\-()]*\d{2}[\s\-()]*\d{2}',
        'Паспорт': r'(?:серия\s+\d{4}\s+номер\s+\d{6})|(?:\b\d{4}\s\d{6}\b|\b\d{2}\s\d{2}\s\d{6}\b)|(?:СЕРИЯ\s+\d{4}\s+НОМЕР\s+\d{6})',
        'ИНН': r'\bИНН[:\s-]*\d[\d\s-]{9,13}\b|\bинн[:\s-]*\d[\d\s-]{9,13}\b|\b\d{10}\b|\b\d{12}\b|(?:ИНН[:\s]*)?\b\d{10}\b|\b\d{12}\b',
        'СНИЛС': r'\d{3}[-\s]?\d{3}[-\s]?\d{3}\s?\d{2}|\b(?:\d{3}[\s\-–—]*\d{3}[\s\-–—]*\d{3}[\s\-–—]*\d{2})\b',
        'СНИЛС1':r"(?:сн[ил]л?с|chнjc|cнилс|CHИJC)?[:\s\-]*\d{3}[\s\-\–—]?\d{3}[\s\-\–—]?\d{3}[\s\-\–—]?\d{2}",


        'ДатаРожд': r'''
            (?:дата\s+рождения[:\s]*)?
            \d{1,2}
            \s+
            (?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)
            \s+
            \d{4}
            \s*г?\.?
        ''',
        'Адресс_П': r'''
            (?:адрес\s+места\s+жительства[:\s]*)?(?:обл\.?\s*[А-ЯЁа-яё]+,\s*)?(?:г\.?\s*[А-ЯЁа-яё]+,\s*)?(?:ул\.?\s*[А-ЯЁа-яё\d\s.-]+,\s*)?(?:д\.?\s*\d+[\w/-]*,\s*)?(?:\d{6})?
        ''',
        'Адрес': r'\b(?:улица|ул|проспект|пр-т|переулок|пер|шоссе|ш|площадь|пл|набережная|наб)\.?\s+[А-ЯЁа-яё\d\s.,-]+\b,\s*(?:дом|д)\.?\s*\d+[\w/]*(?:\s*,\s*(?:кв|квартира|корп|корпус|стр|строение)[\w\d\s./-]*)?'
    }

    for data_type, pattern in patterns.items():
        flags = re.IGNORECASE if data_type == 'ADDRESS' or data_type == 'ADDRESS_FULL' else 0
        matches = re.findall(pattern, text, flags=re.IGNORECASE | re.VERBOSE)
        for match in set(matches):
            if match == '':
                continue
            all_findings.append({
                "type": data_type,
                "value": re.sub(r'\s+', ' ', match).strip(),
                "category": GENERAL_DATA_CLASSIFIER.get(data_type, "Общие")
            })

    #Поиск ОБЩИХ ПДн (Natasha NER)
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_ner(ner_tagger)
    for span in doc.spans:
        if span.type in ['PER']:
            data_type = 'FULL_NAME_NER' if span.type == 'PER' else 'LOCATION_NER'
            span.normalize(morph_vocab)
            if data_type == 'FULL_NAME_NER':
                data_type = 'ФИО'
            else: data_type = 'Расположение'
            all_findings.append({
                "type": data_type, "value": span.normal,
                "category": GENERAL_DATA_CLASSIFIER.get(data_type, "Общие")
            })

    #Поиск СПЕЦИАЛЬНЫХ ПДн
    text_lower = text.lower()
    processed_indices = set()
    for category_name, markers in SPECIAL_DATA_MARKERS.items():
        for marker in markers:
            for match in re.finditer(r'\b' + re.escape(marker) + r'\b', text_lower):
                start, end = match.start(), match.end()
                if any(i in processed_indices for i in range(start, end)):
                    continue

                context_start = max(0, start - 50)
                context_end = min(len(text), end + 100)
                context_snippet = text[context_start:context_end].strip().replace('\n', ' ')

                all_findings.append({
                    "type": category_name,
                    "value": f"Контекст: ...{context_snippet}...",
                    "category": "Специальные"
                })
                processed_indices.update(range(context_start, context_end))

    try:
        result_rubert = ner_pipeline(text)
        for entity in result_rubert:
            if entity['score'] >= 0.8 and len(entity['word']) > 1:
                all_findings.append({
                    "type": entity['entity_group'],
                    "value": entity['word'],
                    "category": "Общие"
                })
    except RuntimeError:
        pass


    text_razd = text.replace("\n", " ").split(" ")
    text_razd = [x for x in text_razd if x]
    for i in range(len(text_razd) - 1):
        if re.search("([0-9]{2}.[0-9]{2}.[0-9]{4})", text_razd[i + 1]) or text_razd[i] in ["Рождения:", "рождения:", "рождения", "Рождения"]:
            all_findings.append({
                "type": "ДатаРожд",
                "value": text_razd[i + 1],
                "category": "Общие"
            })
        if text_razd[i] in all_FIO:
            all_findings.append({
                "type": "ФИО",
                "value": text_razd[i + 1],
                "category": "Общие"})
        if text_razd[i] == "выдан" and len(text_razd[i + 1])>9:
            all_findings.append({
                "type": "Дата",
                "value": text_razd[i + 1],
                "category": "Общие"})
        if text_razd[i] in ["район" ,"района", "Район", "района","р-не",'Р-НЕ','РЕСПУБЛИКЕ','Республике','Р-НА','р-на','P-HA','P-HA']:
            all_findings.append({
                "type": "Адрес1",
                "value": text_razd[i - 1],
                "category": "Общие"})
        if text_razd[i] in ['ГОРОДА','г.', 'города', 'Города', 'квартиры', 'Квартиры', 'кв.', 'КВ.', 'ул.', 'ул', 'Улица', 'улица', 'д.', 'дом', 'Дом','ГОР.','С.','Село','с.','село','ГОРОДЕ','городе', 'ФОР.']:
            all_findings.append({
                "type": "Адрес1",
                "value": text_razd[i + 1],
                "category": "Общие"})
        if text_razd[i] in "№ номер":
            if text_razd[i - 1] in ["Анализы", "анализы", "карта", "Карта пациента", "Пациента"]:
                all_findings.append({
                    "type": "Медицинские данные",
                    "value": text_razd[i + 1],
                    "category": "Специальные"})
    if not all_findings:
        return []
    unique_findings = [dict(t) for t in {tuple(d.items()) for d in all_findings}]
    return unique_findings


def get_file_owner(file_path: str) -> str:
    try:
        sd = win32security.GetFileSecurity(
            file_path,
            win32security.OWNER_SECURITY_INFORMATION
        )
        owner_sid = sd.GetSecurityDescriptorOwner()
        name, domain, _ = win32security.LookupAccountSid(None, owner_sid)
        return f"{domain}\\{name}"
    except Exception:
        return "—"

def load_keywords_from_file(filepath: str) -> List[str]:
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        keywords = [w.strip().lower() for w in content.split(";") if w.strip()]
        return keywords
    except Exception as e:
        return []

def find_keywords_in_text(text: str, keywords: List[str]):
    findings = []
    text_lower = text.lower()

    for kw in keywords:
        for match in re.finditer(r'\b' + re.escape(kw) + r'\b', text_lower):
            context_start = max(0, match.start() - 40)
            context_end = min(len(text), match.end() + 60)

            findings.append({
                "type": "KEYWORD",
                "value": kw,
                "category": "Пользовательские",
                "context": text[context_start:context_end].replace("\n", " ")
            })

    return findings


def extract_text_from_txt(filepath: str):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='cp1251') as f:
                return f.read()
        except Exception as e:
            return None
    except Exception as e:
        return None

def extract_text_from_docx(filepath: str):
    try:
        doc = docx.Document(filepath)
        full_text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        return None


def extract_text_from_pdf(filepath: str):
    try:
        with pdfplumber.open(filepath) as pdf:
            full_text = [page.extract_text() for page in pdf.pages if page.extract_text()]
        return '\n'.join(full_text)
    except Exception as e:
        return None

def faceless_pdf(item):
    filepath = item['file']
    pdf = fitz.open(filepath)
    masked_images = []
    for page_index in range(len(pdf)):
        page = pdf[page_index]
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)

        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        #OCR Paddle
        temp_path = f"{filepath}_temp_page_{page_index}.png"
        img.save(temp_path)

        result = ocr.predict(temp_path)
        rec_texts = result[0]['rec_texts']
        rec_boxes = result[0]['rec_boxes']

        full_text = ' '.join(rec_texts)
        os.remove(temp_path)
        sensitive_values = find_personal_data_in_text(full_text)
        draw = ImageDraw.Draw(img)

        for i, txt in enumerate(rec_texts):
            if not txt.strip():
                continue
            if any(s['value'] in txt for s in sensitive_values):
                box = rec_boxes[i]
                if isinstance(box[0], (int, float, np.integer)):
                    if len(box) == 8:
                        pts = [(box[j], box[j+1]) for j in range(0, 8, 2)]
                    elif len(box) == 4:
                        x1_, y1_, x2_, y2_ = box
                        pts = [(x1_, y1_), (x2_, y1_), (x2_, y2_), (x1_, y2_)]
                    else:
                        continue
                else:
                    pts = box

                xs = [p[0] for p in pts]
                ys = [p[1] for p in pts]
                x1, y1 = min(xs), min(ys)
                x2, y2 = max(xs), max(ys)

                draw.rectangle([x1, y1, x2, y2], fill="black")

        masked_images.append(img)

    pdf.close()
    output_path = f"{filepath}.masked.pdf"

    masked_images[0].save(
        output_path,
        save_all=True,
        append_images=masked_images[1:],
        resolution=300
    )

def secure_delete_file(file_path: str, passes: int = 3):
    if not os.path.isfile(file_path):
        return False
    try:
        length = os.path.getsize(file_path)

        with open(file_path, "r+b", buffering=0) as f:
            for i in range(passes):
                f.seek(0)
                if i % 2 == 0:
                    # случайные данные
                    f.write(os.urandom(length))
                else:
                    # нули
                    f.write(b'\x00' * length)

                f.flush()
                os.fsync(f.fileno())

        success = False

        try:
            os.remove(file_path)
            success = True
        except Exception:
            success = False

        log_destruction(
            file_path=file_path,
            method="Overwrite+Delete",
            passes=passes,
            success=success
        )

        return success

    except Exception as e:
        print(f"[ОШИБКА ЗАТИРАНИЯ] {file_path}: {e}")
        return False

LOG_FILE = "destruction_log.csv"

def log_destruction(
    file_path: str,
    method: str,
    passes: int,
    success: bool
):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    user = getpass.getuser()

    try:
        size = os.path.getsize(file_path)
    except Exception:
        size = "-"

    status = "SUCCESS" if success else "FAILED"

    line = f"{timestamp};{user};{file_path};{size};{method};{passes};{status}\n"

    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(line)


def extract_text_from_odt(filepath: str):
    try:
        doc = load(filepath)
        all_texts = [teletype.extractText(el) for el in doc.getElementsByType(text.P) + doc.getElementsByType(text.H)]
        return '\n'.join(all_texts)
    except Exception as e:
        print(f"Ошибка чтения файла {filepath}: {e}");
        return None



def get_file_age_str(file_path: str) -> str:
    try:
        created_ts = os.path.getctime(file_path)
        created_dt = datetime.datetime.fromtimestamp(created_ts)
        now = datetime.datetime.now()
        delta = now - created_dt
        years = delta.days // 365
        days = delta.days
        hours = delta.seconds // 3600

        return f"{days} дн. {hours} ч."
    except Exception:
        return "—"

def generate_report(scan_results, output_path="pdn_report.xlsx"):
    rows = []
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    owner = getpass.getuser()

    for item in scan_results:
        file_path = item["file"]
        filename = os.path.basename(file_path)
        folder = os.path.dirname(file_path)

        types_set = sorted(set(f["type"] for f in item["findings"]))
        types_string = ", ".join(types_set)

        rows.append({
            "Файл": filename,
            "Путь": folder,
            "Владелец": owner,
            "Типы данных": types_string,
            "Количество": len(item["findings"]),
            "Дата сканирования": now
        })

    df = pd.DataFrame(rows)

    try:
        df.to_excel(output_path, index=False)
        return output_path
    except Exception as e:
        return None

def faceless_image(item):
    filepath = item['file']
    result = ocr.predict(
        input=filepath)

    result_text = []
    for finding in item["findings"]:
        value = finding["value"]
        if value != "":
            result_text.append(value)

    rec_texts = result[0]['rec_texts']
    print(rec_texts)
    rec_boxes = result[0]['rec_boxes']
    image = Image.open(filepath)
    print(result_text)
    for i in range(len(rec_texts)):
        for j in result_text:
            if j in rec_texts[i] or rec_texts[i] in j:
                draw = ImageDraw.Draw(image)
                draw.rectangle([rec_boxes[i][0], rec_boxes[i][1], rec_boxes[i][2], rec_boxes[i][3]], fill='black')
    extension = Path(filepath).suffix
    output_path = f"{filepath}.masked{extension}"
    image.save(output_path)
    return

def archive_files_with_zip(file_paths: List[str], output_zip: str, password: str) -> None:
    with pyzipper.AESZipFile(
        output_zip, 'w', compression=pyzipper.ZIP_DEFLATED, encryption=pyzipper.WZ_AES
    ) as zipf:
        zipf.pwd = password.encode('utf-8')
        for file in file_paths:
            zipf.write(file)

def faceless_documents(item):
    filepath = item['file']
    extension = Path(filepath).suffix.lower()
    sensitive_values = [finding["value"] for finding in item["findings"]]

    #Замена пдн в тексте
    def mask_text(text_m):
        masked = text_m
        for val in sensitive_values:
            if not val.strip():
                continue
            pattern = re.escape(val)
            masked = re.sub(pattern, "***", masked, flags=re.IGNORECASE)
        return masked

    #TXT
    if extension == ".txt":
        text_txt = extract_text_from_txt(filepath)
        if not text_txt:
            return
        masked_text = mask_text(text_txt)

        out_path = f"{filepath}.masked.txt"
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(masked_text)
        print(f"Обезличено - {out_path}")
        return

    #DOCX
    if extension == ".docx":
        try:
            doc = docx.Document(filepath)

            for p in doc.paragraphs:
                p.text = mask_text(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = mask_text(cell.text)

            out_path = f"{filepath}.masked.docx"
            doc.save(out_path)
            print(f"Обезличено - {out_path}")
        except Exception as e:
            print(f"Ошибка при обезличивании DOCX: {e}")
        return

    if extension == ".pdf":
        faceless_pdf(item)
        return

    if extension == ".odt":
        try:
            text_content = extract_text_from_odt(filepath)

            if not text_content:
                return
            masked_text = mask_text(text_content)
            out_path = f"{filepath}.masked.txt"
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(masked_text)

        except Exception as e:
            print(f"Ошибка ODT: {e}")
        return


def faceless_text(final_result):
    img_exts = {'.jpg', '.jpeg', '.png', '.bmp'}
    for item in final_result:
        filename = item['file']
        ext = os.path.splitext(filename)[1].lower()
        if ext in img_exts:
            faceless_image(item)
        else:
            faceless_documents(item)

    return



def extract_text_from_image(filepath: str):
    result = ocr.predict(
        input=filepath)
    rec_texts = result[0]['rec_texts']
    print(rec_texts)
    digit_text = ""
    str_rec_text = ""
    for i in rec_texts:
        if (len(i)<2 and "." not in i) or i =="":
            continue
        i = i.replace('N', "Н")
        if i.isdigit():
            digit_text += i + " "
        str_rec_text += i + " "
    str_rec_text += digit_text
    print(str_rec_text)
    return str_rec_text


def main_scanner(directory_path: str, keywords: List[str] = None, mode: str = "pdn"):
    init_destruction_log()
    all_results = []
    print(mode)
    supported_formats = {'.txt': extract_text_from_txt, '.docx': extract_text_from_docx, '.pdf': extract_text_from_pdf,
                         '.odt': extract_text_from_odt, '.jpg': extract_text_from_image,
                        '.jpeg': extract_text_from_image,'.png': extract_text_from_image}
    for dirpath, _, filenames in os.walk(directory_path):
        for filename in filenames:
            file_path = os.path.join(dirpath, filename)
            file_ext = os.path.splitext(filename)[1].lower()
            findings = []
            if file_ext in supported_formats:
                extractor_func = supported_formats[file_ext]
                text_content = extractor_func(file_path)
                if mode in ("pdn", "pdn+keywords"):
                    if text_content and text_content.strip():
                        findings.extend(find_personal_data_in_text(text_content))

                if mode in ("keywords", "pdn+keywords") and keywords:
                    findings.extend(find_keywords_in_text(text_content, keywords))

                if findings:
                    file_time_creat = get_file_age_str(file_path)
                    all_results.append({'file': file_path, 'findings': findings, 'file_time_creat': file_time_creat})
                    for finding in findings:
                        print(
                            f"  -> Категория: {finding['category']}, Тип: {finding['type']}, Значение: '{finding['value']}'")
                else:
                    print("Персональные данные не найдены.")

            else:
                print(" Не удалось извлечь текст или файл пустой.")

    return all_results




if __name__ == "__main__":
    TARGET_DIRECTORY = "C:/Users/vova/Desktop/ПО/Доки/Смесь"

    if os.path.isdir(TARGET_DIRECTORY):
        final_results = main_scanner(TARGET_DIRECTORY)
        dir_save = []
        for item in final_results:
            dir_save.append(item['file'])
        answer = input("")
        if answer.lower() == "да":
            faceless_text(final_results)
        answer = input("")
        if answer.lower() == "да":
            archive_files_with_zip(dir_save, 'secure_archive.zip', 'mypassword')
        else:
            generate_report(final_results)


